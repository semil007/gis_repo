"""
Enhanced file upload validation with comprehensive checks and visual feedback.
Provides detailed validation for file format, size, content, and security.
"""

import streamlit as st
import os
import hashlib
from typing import Dict, List, Optional, Tuple
from datetime import datetime
import tempfile
import PyPDF2
from docx import Document

# Optional import for MIME type detection
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False


class UploadValidator:
    """Comprehensive file upload validator with security and content checks."""
    
    def __init__(self):
        self.max_file_size_mb = 100
        self.max_file_size_bytes = self.max_file_size_mb * 1024 * 1024
        self.supported_mime_types = {
            'application/pdf': '.pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx'
        }
        self.supported_extensions = ['.pdf', '.docx']
        
    def validate_comprehensive(self, uploaded_file) -> Dict[str, any]:
        """
        Perform comprehensive validation of uploaded file.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Comprehensive validation results dictionary
        """
        validation_result = {
            'is_valid': True,
            'errors': [],
            'warnings': [],
            'info': [],
            'file_info': {},
            'security_checks': {},
            'content_analysis': {},
            'recommendations': []
        }
        
        if uploaded_file is None:
            validation_result['is_valid'] = False
            validation_result['errors'].append("No file uploaded")
            return validation_result
            
        try:
            # Basic file information
            file_info = self._extract_file_info(uploaded_file)
            validation_result['file_info'] = file_info
            
            # Size validation
            size_check = self._validate_file_size(file_info['size'])
            validation_result['errors'].extend(size_check['errors'])
            validation_result['warnings'].extend(size_check['warnings'])
            
            # Format validation
            format_check = self._validate_file_format(uploaded_file, file_info)
            validation_result['errors'].extend(format_check['errors'])
            validation_result['warnings'].extend(format_check['warnings'])
            validation_result['security_checks'].update(format_check['security_checks'])
            
            # Content validation
            content_check = self._validate_file_content(uploaded_file)
            validation_result['content_analysis'] = content_check
            validation_result['warnings'].extend(content_check.get('warnings', []))
            validation_result['info'].extend(content_check.get('info', []))
            
            # Generate recommendations
            validation_result['recommendations'] = self._generate_recommendations(validation_result)
            
            # Set overall validity
            validation_result['is_valid'] = len(validation_result['errors']) == 0
            
        except Exception as e:
            validation_result['is_valid'] = False
            validation_result['errors'].append(f"Validation error: {str(e)}")
            
        return validation_result
        
    def _extract_file_info(self, uploaded_file) -> Dict[str, any]:
        """Extract basic file information."""
        file_name = uploaded_file.name
        file_size = uploaded_file.size
        file_extension = os.path.splitext(file_name)[1].lower()
        
        # Calculate file hash for integrity checking
        file_content = uploaded_file.getvalue()
        file_hash = hashlib.md5(file_content).hexdigest()
        
        return {
            'name': file_name,
            'size': file_size,
            'size_mb': file_size / (1024 * 1024),
            'extension': file_extension,
            'hash': file_hash,
            'upload_time': datetime.now(),
            'content_length': len(file_content)
        }
        
    def _validate_file_size(self, file_size: int) -> Dict[str, List[str]]:
        """Validate file size against limits."""
        result = {'errors': [], 'warnings': []}
        
        if file_size > self.max_file_size_bytes:
            result['errors'].append(
                f"File size ({file_size / (1024 * 1024):.1f}MB) exceeds maximum allowed size ({self.max_file_size_mb}MB)"
            )
        elif file_size > (self.max_file_size_bytes * 0.8):
            result['warnings'].append(
                f"Large file detected ({file_size / (1024 * 1024):.1f}MB). Processing may take longer than usual."
            )
        elif file_size < 1024:  # Less than 1KB
            result['warnings'].append(
                "Very small file detected. Please ensure the file contains the expected content."
            )
            
        return result
        
    def _validate_file_format(self, uploaded_file, file_info: Dict) -> Dict[str, any]:
        """Validate file format using multiple methods."""
        result = {
            'errors': [],
            'warnings': [],
            'security_checks': {}
        }
        
        file_extension = file_info['extension']
        
        # Extension validation
        if file_extension not in self.supported_extensions:
            result['errors'].append(
                f"Unsupported file format '{file_extension}'. Supported formats: {', '.join(self.supported_extensions)}"
            )
            return result
            
        # MIME type validation (if python-magic is available)
        if MAGIC_AVAILABLE:
            try:
                file_content = uploaded_file.getvalue()
                mime_type = magic.from_buffer(file_content, mime=True)
                result['security_checks']['detected_mime_type'] = mime_type
                
                if mime_type not in self.supported_mime_types:
                    result['warnings'].append(
                        f"File extension '{file_extension}' doesn't match detected MIME type '{mime_type}'"
                    )
            except Exception as e:
                result['warnings'].append(f"Could not verify MIME type: {str(e)}")
        else:
            result['security_checks']['mime_check'] = 'python-magic not available'
            
        # File signature validation
        signature_check = self._validate_file_signature(uploaded_file, file_extension)
        result['security_checks'].update(signature_check)
        
        return result
        
    def _validate_file_signature(self, uploaded_file, expected_extension: str) -> Dict[str, any]:
        """Validate file signature (magic bytes)."""
        result = {'signature_valid': False, 'detected_format': None}
        
        try:
            file_content = uploaded_file.getvalue()
            
            # PDF signature
            if file_content.startswith(b'%PDF'):
                result['signature_valid'] = expected_extension == '.pdf'
                result['detected_format'] = 'PDF'
                
            # DOCX signature (ZIP-based)
            elif file_content.startswith(b'PK\x03\x04'):
                result['signature_valid'] = expected_extension == '.docx'
                result['detected_format'] = 'ZIP-based (likely DOCX)'
                
            else:
                result['signature_valid'] = False
                result['detected_format'] = 'Unknown'
                
        except Exception as e:
            result['error'] = str(e)
            
        return result
        
    def _validate_file_content(self, uploaded_file) -> Dict[str, any]:
        """Validate file content and extract basic information."""
        result = {
            'content_valid': False,
            'page_count': 0,
            'has_text': False,
            'estimated_records': 0,
            'warnings': [],
            'info': []
        }
        
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        
        try:
            if file_extension == '.pdf':
                result.update(self._analyze_pdf_content(uploaded_file))
            elif file_extension == '.docx':
                result.update(self._analyze_docx_content(uploaded_file))
                
        except Exception as e:
            result['warnings'].append(f"Content analysis failed: {str(e)}")
            
        return result
        
    def _analyze_pdf_content(self, uploaded_file) -> Dict[str, any]:
        """Analyze PDF content for validation."""
        result = {
            'content_valid': False,
            'page_count': 0,
            'has_text': False,
            'estimated_records': 0,
            'info': [],
            'warnings': []
        }
        
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            result['page_count'] = len(pdf_reader.pages)
            result['info'].append(f"PDF contains {result['page_count']} page(s)")
            
            # Extract text from first few pages to check content
            text_content = ""
            pages_to_check = min(3, len(pdf_reader.pages))
            
            for i in range(pages_to_check):
                try:
                    page_text = pdf_reader.pages[i].extract_text()
                    text_content += page_text
                except Exception as e:
                    result['warnings'].append(f"Could not extract text from page {i+1}: {str(e)}")
                    
            result['has_text'] = len(text_content.strip()) > 0
            
            if result['has_text']:
                result['content_valid'] = True
                result['info'].append(f"Extracted {len(text_content)} characters of text")
                
                # Estimate number of records based on common HMO keywords
                hmo_keywords = ['licence', 'license', 'hmo', 'occupancy', 'manager', 'holder']
                keyword_count = sum(text_content.lower().count(keyword) for keyword in hmo_keywords)
                result['estimated_records'] = max(1, keyword_count // 5)  # Rough estimate
                result['info'].append(f"Estimated {result['estimated_records']} potential record(s)")
            else:
                result['warnings'].append("No readable text found. Document may be scanned or corrupted.")
                
        except Exception as e:
            result['warnings'].append(f"PDF analysis failed: {str(e)}")
            
        return result
        
    def _analyze_docx_content(self, uploaded_file) -> Dict[str, any]:
        """Analyze DOCX content for validation."""
        result = {
            'content_valid': False,
            'page_count': 1,  # DOCX doesn't have fixed pages
            'has_text': False,
            'estimated_records': 0,
            'info': [],
            'warnings': []
        }
        
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            
            # Read DOCX
            doc = Document(uploaded_file)
            
            # Extract text content
            text_content = ""
            paragraph_count = 0
            table_count = len(doc.tables)
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
                    paragraph_count += 1
                    
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                        
            result['has_text'] = len(text_content.strip()) > 0
            result['info'].append(f"Document contains {paragraph_count} paragraph(s) and {table_count} table(s)")
            
            if result['has_text']:
                result['content_valid'] = True
                result['info'].append(f"Extracted {len(text_content)} characters of text")
                
                # Estimate number of records
                hmo_keywords = ['licence', 'license', 'hmo', 'occupancy', 'manager', 'holder']
                keyword_count = sum(text_content.lower().count(keyword) for keyword in hmo_keywords)
                result['estimated_records'] = max(1, keyword_count // 5)
                result['info'].append(f"Estimated {result['estimated_records']} potential record(s)")
            else:
                result['warnings'].append("No readable text found in document.")
                
        except Exception as e:
            result['warnings'].append(f"DOCX analysis failed: {str(e)}")
            
        return result
        
    def _generate_recommendations(self, validation_result: Dict) -> List[str]:
        """Generate recommendations based on validation results."""
        recommendations = []
        
        file_info = validation_result.get('file_info', {})
        content_analysis = validation_result.get('content_analysis', {})
        
        # Size-based recommendations
        if file_info.get('size_mb', 0) > 50:
            recommendations.append("Consider splitting large documents for faster processing")
            
        # Content-based recommendations
        if not content_analysis.get('has_text', False):
            recommendations.append("Document appears to be scanned. OCR processing will be used, which may take longer")
            
        if content_analysis.get('estimated_records', 0) == 0:
            recommendations.append("No HMO-related content detected. Please verify this is the correct document type")
            
        # Format-based recommendations
        if file_info.get('extension') == '.pdf' and content_analysis.get('page_count', 0) > 20:
            recommendations.append("Large PDF detected. Consider using DOCX format for better processing speed")
            
        return recommendations


class VisualFeedback:
    """Provides visual feedback for upload operations."""
    
    @staticmethod
    def show_upload_zone_feedback(is_dragover: bool = False) -> None:
        """Show visual feedback for drag-and-drop zone."""
        if is_dragover:
            st.markdown("""
            <div style="
                border: 3px dashed #1f77b4;
                border-radius: 10px;
                padding: 40px;
                text-align: center;
                background-color: #e6f3ff;
                color: #1f77b4;
                font-size: 18px;
                font-weight: bold;
            ">
                📁 Drop your file here!
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div style="
                border: 2px dashed #cccccc;
                border-radius: 10px;
                padding: 40px;
                text-align: center;
                background-color: #f9f9f9;
                color: #666666;
                font-size: 16px;
            ">
                📁 Drag and drop your file here or click to browse
            </div>
            """, unsafe_allow_html=True)
            
    @staticmethod
    def show_validation_summary(validation_result: Dict) -> None:
        """Display comprehensive validation summary."""
        if validation_result['is_valid']:
            st.success("✅ File validation passed!")
            
            # Show file information
            file_info = validation_result['file_info']
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("File Size", f"{file_info['size_mb']:.1f} MB")
            with col2:
                st.metric("Format", file_info['extension'].upper())
            with col3:
                content = validation_result['content_analysis']
                st.metric("Est. Records", content.get('estimated_records', 0))
                
            # Show recommendations
            if validation_result['recommendations']:
                with st.expander("💡 Recommendations", expanded=False):
                    for rec in validation_result['recommendations']:
                        st.info(f"• {rec}")
                        
        else:
            st.error("❌ File validation failed!")
            for error in validation_result['errors']:
                st.error(f"• {error}")
                
        # Show warnings
        if validation_result['warnings']:
            for warning in validation_result['warnings']:
                st.warning(f"⚠️ {warning}")
                
        # Show info messages
        if validation_result['info']:
            with st.expander("ℹ️ File Analysis Details", expanded=False):
                for info in validation_result['info']:
                    st.info(f"• {info}")


class UploadProgressIndicator:
    """Enhanced upload progress indicator with detailed feedback."""
    
    def __init__(self):
        self.steps = [
            "Receiving file...",
            "Validating file format...",
            "Checking file size...",
            "Analyzing content...",
            "Running security checks...",
            "Saving to temporary storage...",
            "Upload completed!"
        ]
        
    def show_progress(self, step: int, custom_message: str = None) -> None:
        """Show upload progress with detailed steps."""
        progress = min(step / len(self.steps), 1.0)
        message = custom_message or (self.steps[step] if step < len(self.steps) else "Complete")
        
        st.progress(progress)
        st.text(f"Step {step + 1}/{len(self.steps)}: {message}")
        
        if progress >= 1.0:
            st.success("✅ Upload completed successfully!")
            
    def show_error(self, error_message: str, step: int = None) -> None:
        """Show upload error with context."""
        st.error(f"❌ Upload failed: {error_message}")
        
        if step is not None and step < len(self.steps):
            st.error(f"Failed at step {step + 1}: {self.steps[step]}")
            
        # Show retry option
        if st.button("🔄 Try Again", use_container_width=True):
            st.rerun()