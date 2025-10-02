"""
DOCX document processor implementation using python-docx.
"""

import re
from pathlib import Path
from typing import Dict, List, Optional, Union

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from loguru import logger

from .base import DocumentProcessor, ProcessingResult, ProcessingStatus, DocumentProcessingError


class DOCXProcessingError(DocumentProcessingError):
    """DOCX-specific processing error."""
    pass


class DOCXProcessor(DocumentProcessor):
    """
    DOCX document processor using python-docx for text extraction,
    table processing, and metadata extraction.
    """
    
    def __init__(self, config: Optional[Dict] = None):
        """
        Initialize DOCX processor with configuration.
        
        Args:
            config: Configuration dictionary with processing options
        """
        super().__init__(config)
        
        # DOCX-specific configuration
        self.preserve_formatting = self.config.get("preserve_formatting", True)
        self.extract_headers_footers = self.config.get("extract_headers_footers", True)
        self.table_detection = self.config.get("table_detection", True)
        self.max_tables = self.config.get("max_tables", 50)
        
    def process_docx(self, file_path: Union[str, Path]) -> ProcessingResult:
        """
        Process DOCX document with text extraction, table processing, and metadata extraction.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ProcessingResult with extracted data
        """
        file_path = Path(file_path)
        
        try:
            logger.info(f"Starting DOCX processing: {file_path}")
            
            # Initialize result containers
            extracted_text = ""
            extracted_data = {}
            confidence_scores = {}
            flagged_fields = []
            error_messages = []
            processing_metadata = {
                "file_path": str(file_path),
                "processor": "DOCXProcessor"
            }
            
            # Open and analyze DOCX
            document = Document(file_path)
            
            # Extract document metadata
            metadata = self._extract_metadata(document)
            if metadata:
                extracted_data["metadata"] = metadata
                processing_metadata.update(metadata)
                
            # Extract text content while preserving structure
            text_content = self._extract_text_content(document)
            extracted_text = text_content["full_text"]
            
            # Store structured text components
            if text_content.get("paragraphs"):
                extracted_data["paragraphs"] = text_content["paragraphs"]
                
            if text_content.get("headings"):
                extracted_data["headings"] = text_content["headings"]
                
            # Extract tables if enabled
            if self.table_detection:
                tables_data = self._extract_tables(document)
                if tables_data:
                    extracted_data["tables"] = tables_data
                    processing_metadata["tables_found"] = len(tables_data)
                    
            # Extract headers and footers if enabled
            if self.extract_headers_footers:
                headers_footers = self._extract_headers_footers(document)
                if headers_footers:
                    extracted_data["headers_footers"] = headers_footers
                    
            # Basic data extraction from text
            basic_data = self._extract_basic_data(extracted_text)
            extracted_data.update(basic_data)
            
            # Calculate confidence scores
            confidence_scores = self._calculate_confidence_scores(extracted_data, extracted_text)
            
            # Flag low confidence fields
            confidence_threshold = self.config.get("confidence_threshold", 0.7)
            flagged_fields = [
                field for field, score in confidence_scores.items() 
                if score < confidence_threshold
            ]
            
            # Determine processing status
            if error_messages:
                status = ProcessingStatus.PARTIAL if extracted_text else ProcessingStatus.FAILED
            else:
                status = ProcessingStatus.COMPLETED
                
            logger.info(f"DOCX processing completed: {len(extracted_text)} chars, {len(extracted_data)} fields")
            
            return ProcessingResult(
                extracted_text=extracted_text,
                extracted_data=extracted_data,
                confidence_scores=confidence_scores,
                flagged_fields=flagged_fields,
                processing_metadata=processing_metadata,
                status=status,
                error_messages=error_messages,
                page_count=None,  # DOCX doesn't have fixed pages
                tables_found=len(tables_data) if self.table_detection and 'tables' in extracted_data else 0,
                ocr_used=False
            )
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {str(e)}")
            return ProcessingResult(
                extracted_text="",
                extracted_data={},
                confidence_scores={},
                flagged_fields=[],
                processing_metadata={"error": str(e), "file_path": str(file_path)},
                status=ProcessingStatus.FAILED,
                error_messages=[str(e)],
                ocr_used=False
            )
            
    def _extract_metadata(self, document: Document) -> Dict:
        """
        Extract document metadata and properties.
        
        Args:
            document: python-docx Document object
            
        Returns:
            Dictionary with document metadata
        """
        metadata = {}
        
        try:
            core_props = document.core_properties
            
            # Basic document properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()
            if core_props.last_modified_by:
                metadata["last_modified_by"] = core_props.last_modified_by
                
            # Document statistics
            paragraph_count = len(document.paragraphs)
            table_count = len(document.tables)
            
            metadata["paragraph_count"] = paragraph_count
            metadata["table_count"] = table_count
            
            logger.debug(f"Extracted metadata: {len(metadata)} properties")
            
        except Exception as e:
            logger.warning(f"Metadata extraction failed: {str(e)}")
            
        return metadata
        
    def _extract_text_content(self, document: Document) -> Dict:
        """
        Extract text content while preserving document structure.
        
        Args:
            document: python-docx Document object
            
        Returns:
            Dictionary with structured text content
        """
        content = {
            "full_text": "",
            "paragraphs": [],
            "headings": []
        }
        
        try:
            text_parts = []
            
            for paragraph in document.paragraphs:
                para_text = paragraph.text.strip()
                
                if para_text:
                    # Check if paragraph is a heading
                    if paragraph.style.name.startswith('Heading'):
                        content["headings"].append({
                            "text": para_text,
                            "level": paragraph.style.name,
                            "position": len(content["paragraphs"])
                        })
                        
                    # Store paragraph with formatting info if enabled
                    if self.preserve_formatting:
                        para_info = {
                            "text": para_text,
                            "style": paragraph.style.name,
                            "is_heading": paragraph.style.name.startswith('Heading')
                        }
                        
                        # Extract run-level formatting
                        runs_info = []
                        for run in paragraph.runs:
                            if run.text.strip():
                                run_info = {
                                    "text": run.text,
                                    "bold": run.bold,
                                    "italic": run.italic,
                                    "underline": run.underline
                                }
                                runs_info.append(run_info)
                                
                        if runs_info:
                            para_info["runs"] = runs_info
                            
                        content["paragraphs"].append(para_info)
                    else:
                        content["paragraphs"].append({"text": para_text})
                        
                    text_parts.append(para_text)
                    
            content["full_text"] = "\n\n".join(text_parts)
            
            logger.debug(f"Extracted {len(content['paragraphs'])} paragraphs, {len(content['headings'])} headings")
            
        except Exception as e:
            logger.error(f"Text content extraction failed: {str(e)}")
            
        return content
        
    def _extract_tables(self, document: Document) -> List[Dict]:
        """
        Extract and process tables from DOCX document.
        
        Args:
            document: python-docx Document object
            
        Returns:
            List of table data dictionaries
        """
        tables_data = []
        
        try:
            table_count = min(len(document.tables), self.max_tables)
            
            for i, table in enumerate(document.tables[:table_count]):
                table_data = self._process_table(table, i)
                if table_data:
                    tables_data.append(table_data)
                    
            logger.debug(f"Extracted {len(tables_data)} tables")
            
        except Exception as e:
            logger.error(f"Table extraction failed: {str(e)}")
            
        return tables_data
        
    def _process_table(self, table: Table, table_id: int) -> Optional[Dict]:
        """
        Process individual table and extract data.
        
        Args:
            table: python-docx Table object
            table_id: Unique identifier for the table
            
        Returns:
            Dictionary with table data or None if processing fails
        """
        try:
            rows_data = []
            headers = []
            
            # Process table rows
            for row_idx, row in enumerate(table.rows):
                row_data = []
                
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                    
                if row_idx == 0 and self._looks_like_header(row_data):
                    headers = row_data
                else:
                    rows_data.append(row_data)
                    
            # Skip empty tables
            if not rows_data and not headers:
                return None
                
            # Create table dictionary
            table_dict = {
                "table_id": table_id,
                "headers": headers,
                "rows": rows_data,
                "shape": (len(rows_data) + (1 if headers else 0), len(headers) if headers else len(rows_data[0]) if rows_data else 0),
                "extraction_method": "python-docx"
            }
            
            # Convert to records format if headers are available
            if headers and rows_data:
                records = []
                for row in rows_data:
                    # Pad row to match header length
                    padded_row = row + [''] * (len(headers) - len(row))
                    record = dict(zip(headers, padded_row))
                    records.append(record)
                table_dict["data"] = records
                
            logger.debug(f"Processed table {table_id} with shape {table_dict['shape']}")
            
            return table_dict
            
        except Exception as e:
            logger.warning(f"Failed to process table {table_id}: {str(e)}")
            return None
            
    def _looks_like_header(self, row_data: List[str]) -> bool:
        """
        Determine if a row looks like a table header.
        
        Args:
            row_data: List of cell values in the row
            
        Returns:
            True if row appears to be a header
        """
        if not row_data:
            return False
            
        # Check for common header indicators
        header_indicators = [
            lambda x: any(word in x.lower() for word in ['name', 'address', 'date', 'reference', 'licence', 'license']),
            lambda x: x.isupper() and len(x) > 2,
            lambda x: ':' in x or '#' in x,
            lambda x: len(x.split()) <= 3 and x.replace(' ', '').isalpha()
        ]
        
        # Count cells that look like headers
        header_like_count = sum(
            1 for cell in row_data 
            if cell.strip() and any(indicator(cell) for indicator in header_indicators)
        )
        
        # Consider it a header if more than half the cells look like headers
        return header_like_count > len(row_data) / 2
        
    def _extract_headers_footers(self, document: Document) -> Dict:
        """
        Extract headers and footers from document sections.
        
        Args:
            document: python-docx Document object
            
        Returns:
            Dictionary with headers and footers content
        """
        headers_footers = {
            "headers": [],
            "footers": []
        }
        
        try:
            for section in document.sections:
                # Extract headers
                if section.header:
                    header_text = ""
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            header_text += paragraph.text.strip() + "\n"
                    if header_text:
                        headers_footers["headers"].append(header_text.strip())
                        
                # Extract footers
                if section.footer:
                    footer_text = ""
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            footer_text += paragraph.text.strip() + "\n"
                    if footer_text:
                        headers_footers["footers"].append(footer_text.strip())
                        
            logger.debug(f"Extracted {len(headers_footers['headers'])} headers, {len(headers_footers['footers'])} footers")
            
        except Exception as e:
            logger.warning(f"Headers/footers extraction failed: {str(e)}")
            
        return headers_footers if headers_footers["headers"] or headers_footers["footers"] else {}
        
    def _extract_basic_data(self, text: str) -> Dict:
        """
        Extract basic HMO data patterns from text.
        
        Args:
            text: Extracted text content
            
        Returns:
            Dictionary with extracted data fields
        """
        data = {}
        
        try:
            # Extract potential reference numbers
            ref_patterns = [
                r'(?:ref|reference|licence|license)[\s:]+([A-Z0-9\-/]+)',
                r'([A-Z]{2,}\d{4,})',  # Pattern like ABC1234
                r'(\d{4,}/[A-Z0-9]+)'  # Pattern like 2023/ABC123
            ]
            
            references = []
            for pattern in ref_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                references.extend(matches)
                
            if references:
                data["potential_references"] = list(set(references))
                
            # Extract potential addresses (basic pattern)
            address_pattern = r'(\d+\s+[A-Za-z\s]+(?:Street|St|Road|Rd|Avenue|Ave|Lane|Ln|Drive|Dr|Close|Cl|Way|Place|Pl)[^,\n]*(?:,\s*[A-Za-z\s]+)*)'
            addresses = re.findall(address_pattern, text, re.IGNORECASE)
            if addresses:
                data["potential_addresses"] = addresses[:10]  # Limit to first 10
                
            # Extract potential dates
            date_patterns = [
                r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
                r'(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{2,4})',
                r'(\d{4}-\d{2}-\d{2})'
            ]
            
            dates = []
            for pattern in date_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                dates.extend(matches)
                
            if dates:
                data["potential_dates"] = list(set(dates))
                
            # Extract potential numbers (occupancy, etc.)
            number_patterns = [
                r'(?:occupancy|capacity|persons?)[\s:]+(\d+)',
                r'(?:max|maximum)[\s:]+(\d+)',
                r'(\d+)\s+(?:persons?|people|occupants?)'
            ]
            
            numbers = []
            for pattern in number_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                numbers.extend(matches)
                
            if numbers:
                data["potential_numbers"] = [int(n) for n in numbers if n.isdigit()]
                
        except Exception as e:
            logger.error(f"Basic data extraction failed: {str(e)}")
            
        return data
        
    def _calculate_confidence_scores(self, extracted_data: Dict, text: str) -> Dict[str, float]:
        """
        Calculate confidence scores for extracted data.
        
        Args:
            extracted_data: Extracted data dictionary
            text: Source text
            
        Returns:
            Dictionary with confidence scores for each field
        """
        confidence_scores = {}
        
        try:
            # Base confidence for DOCX (higher than PDF/OCR)
            base_confidence = 0.9
            
            # Text quality indicators
            text_length = len(text)
            text_quality = min(1.0, text_length / 1000)  # Normalize by expected length
            
            for field, value in extracted_data.items():
                if field == "tables":
                    # Table confidence based on structure
                    if value and len(value) > 0:
                        # DOCX tables are generally well-structured
                        confidence_scores[field] = 0.95
                    else:
                        confidence_scores[field] = 0.0
                        
                elif field == "metadata":
                    # Metadata is highly reliable in DOCX
                    confidence_scores[field] = 0.98
                    
                elif field in ["paragraphs", "headings", "headers_footers"]:
                    # Structural elements are reliable
                    confidence_scores[field] = 0.95
                    
                elif field.startswith("potential_"):
                    # Pattern-based extraction confidence
                    if isinstance(value, list) and value:
                        # More matches generally indicate higher confidence
                        match_confidence = min(0.9, len(value) * 0.2 + 0.4)
                        confidence_scores[field] = match_confidence * base_confidence * text_quality
                    else:
                        confidence_scores[field] = 0.0
                else:
                    # Default confidence for other fields
                    confidence_scores[field] = base_confidence * text_quality
                    
        except Exception as e:
            logger.error(f"Confidence calculation failed: {str(e)}")
            
        return confidence_scores
        
    def process_pdf(self, file_path: Union[str, Path]) -> ProcessingResult:
        """
        PDF processing not implemented in DOCX processor.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            ProcessingResult indicating unsupported operation
        """
        return ProcessingResult(
            extracted_text="",
            extracted_data={},
            confidence_scores={},
            flagged_fields=[],
            processing_metadata={"error": "PDF processing not supported by DOCXProcessor"},
            status=ProcessingStatus.FAILED,
            error_messages=["PDF processing not supported by DOCXProcessor"],
            ocr_used=False
        )